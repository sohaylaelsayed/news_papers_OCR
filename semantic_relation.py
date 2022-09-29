import operator
import re
import docx
from sentence_transformers import SentenceTransformer, util

model = SentenceTransformer('paraphrase-distilroberta-base-v1')


def remove_tags(text):
    text = re.sub('►', ". ", text)
    text = re.sub('-\n', "", text)
    text = re.sub('\s\n(?=/^[\u0621-\u064A\u0660-\u0669 ]+$/)', " ", text)
    text = re.sub('\n\s(?=/^[\u0621-\u064A\u0660-\u0669 ]+$/)', " ", text)
    text = re.sub('\n(?=/^[\u0621-\u064A\u0660-\u0669 ]+$/{2})', " ", text)
    if re.search("/^[\u0621-\u064A\u0660-\u0669 ]+$/", text):
        i = re.search("/^[\u0621-\u064A\u0660-\u0669 ]+$/", text).start()
        b = re.search("/^[\u0621-\u064A\u0660-\u0669 ]+$/", text).end()
        new = text[i:b].replace("\n", " ")
        text = text[:i] + new + text[b + 1:]

    text = re.sub('\n(?=/^[\u0621-\u064A\u0660-\u0669 ]+$/)', " ", text)
    text = re.sub('\n(?=/^[\u0621-\u064A\u0660-\u0669 ]+$/)', ". ", text)
    text = re.sub('[.]\s(?=/^[\u0621-\u064A\u0660-\u0669 ]+$/)', ".", text)
    text = re.sub('[.](?=/^[\u0621-\u064A\u0660-\u0669 ]+$/)', ". ", text)
    text = re.sub('[.]\n(?=/^[\u0621-\u064A\u0660-\u0669 ]+$/)', ". ", text)
    text = re.sub('\s{1,}', " ", text)
    if re.search('/^[\u0621-\u064A\u0660-\u0669 ]+$/', text):
        i_1 = re.search('/^[\u0621-\u064A\u0660-\u0669 ]+$/', text).start()
        b_1 = re.search('/^[\u0621-\u064A\u0660-\u0669 ]+$/', text).end()
        new_1 = text[i_1:b_1].replace("-", "")
        text = text[:i_1] + new_1 + text[b_1 - 1 + 1:]

    text = re.sub('\n', " ", text)

    return re.sub('\xa0', "", text)


def remove_tags_input(text):
    text = re.sub('-\n', "", text)
    text = re.sub('.\n', ". ", text)
    text = re.sub('[.].\s(?=/^[\u0621-\u064A\u0660-\u0669 ]+$/)', ".", text)
    text = re.sub('\n', "", text)
    return re.sub('\xa0', "", text)


def read_word(path):
    text = []
    doc = docx.Document(path)
    all_paras = doc.paragraphs
    for para in all_paras:

        tex = remove_tags(para.text)

        tex = [i.strip() for i in remove_tags(tex).split(". ")]
        i = 0
        for x in tex:

            if len(x) > 1:
                if x[-1] == ".":
                    pass
                else:
                    x = x + "."
                text.append(x)

    return text


def check_format(path):
    if path.endswith(".docx"):
        return read_word(path)
    else:
        print("The file path is wrong")
        return None


def string_similarity(input_text, text, len_corpus, len_input_text):
    sim_dict = {}

    embeddings1 = model.encode(input_text, convert_to_tensor=False)
    for t in text:
        ind = text.index(t)
        sent2 = t
        for i in range(1, len_input_text):
            pick = ind + i
            if pick >= len_corpus:
                pass
            else:
                get_text = text[pick]
                sent2 = sent2 + ". " + get_text
        embeddings2 = model.encode(sent2, convert_to_tensor=False)
        cosine_scores = util.pytorch_cos_sim(embeddings1, embeddings2)
        sim_dict[ind] = cosine_scores

    sorted_d = sorted(sim_dict.items(), key=operator.itemgetter(1), reverse=True)

    return sorted_d


def check_threshold(sorted_d, threshold):
    new_sorted = {}
    for index in sorted_d:
        ind = index[0]
        sim = index[1].numpy()[0][0]
        if sim >= threshold:
            new_sorted[ind] = sim

    return new_sorted


def create_doc(input_text, sorted_d, text, len_input_text):
    mydoc = docx.Document()
    mydoc.add_heading("String: {}".format(input_text), 2)
    x = 1
    for index, value in sorted_d.items():
        ind = index
        val = value
        out = text[ind]
        for i in range(1, len_input_text):
            out = out + ". " + text[ind + i]

        mydoc.add_paragraph("Match: {}".format(x))
        mydoc.add_paragraph("Similartity: {:.2f}".format(val))
        mydoc.add_paragraph(out)

        mydoc.add_paragraph("---------------------")
        x = x + 1

    mydoc.save("similarity_report.docx")


def string_match(string, path, threshold, format):
    text = check_format(path)

    len_corpus = len(text)
    input_text = remove_tags_input(string)
    ab = input_text.split(". ")
    ab = [i for i in ab if len(i) > 1]
    len_input_text = len(ab)

    sorted_d = string_similarity(input_text, text, len_corpus, len_input_text)

    sorted_d = check_threshold(sorted_d, threshold)
    if format.lower() == "word":
        create_doc(input_text, sorted_d, text, len_input_text)
    else:
        print("The output file format is not correct")


def check_doc_emb(embeddings, embeddings1):
    added = []
    modified = []
    deleted = []
    for i in embeddings:
        ind = embeddings.index(i)
        count_added = 0
        count_modified = 0
        for x in embeddings1:
            cosine_scores = util.pytorch_cos_sim(i, x).numpy()[0][0]
            if cosine_scores >= 0.98:

                added.append(ind)
                count_added = 1
                break
            elif 0.98 > cosine_scores > 0.80 and count_modified == 0:
                saved_index = ind
                count_modified = 1
        if count_added == 0:
            if count_modified == 1:
                modified.append(saved_index)
            else:
                deleted.append(ind)

    return added, modified, deleted


def create_doc_doc(added, modified, deleted, text, text1):
    mydoc = docx.Document()
    mydoc.add_heading("Added Paragraph", 2)
    i = 1
    for index in added:
        out = text1[index]
        mydoc.add_paragraph("Match:{}".format(i))
        mydoc.add_paragraph(out)
        mydoc.add_paragraph("---------------------")
        i = i + 1

    mydoc.add_heading("Modified Paragraph", 2)
    i = 1
    for index in modified:
        out = text1[index]
        mydoc.add_paragraph("Match: {}".format(i))
        mydoc.add_paragraph(out)
        mydoc.add_paragraph("---------------------")
        i = i + 1

    mydoc.add_heading("Deleted Paragraph", 2)
    i = 1
    for index in deleted:
        out = text[index]
        mydoc.add_paragraph("Match: {}".format(i))
        mydoc.add_paragraph(out)
        mydoc.add_paragraph("---------------------")
        i = i + 1

    mydoc.save("similarity_report_doc.docx")